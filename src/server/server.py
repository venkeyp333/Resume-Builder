from flask import Flask, request, jsonify, send_file
from flask_mysqldb import MySQL
from flask_bcrypt import Bcrypt
from flask_cors import CORS
from docx import Document
import base64

from io import BytesIO

# import io

import os

app = Flask(__name__)
bcrypt = Bcrypt(app)
CORS(app)

# MySQL configuration
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'  # Replace with your MySQL username
app.config['MYSQL_PASSWORD'] = 'cdac'  # Replace with your MySQL password
app.config['MYSQL_DB'] = 'user_auth'

mysql = MySQL(app)



# Function to check allowed file extensions (for image files)
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# Route for user registration
@app.route('/register', methods=['POST'])
def register():
    data = request.json
    username = data.get('username')
    email = data.get('email')
    password = data.get('password')
    roles = data.get('roles')


    # Check for missing fields
    if not username or not email or not password:
        return jsonify({'message': 'All fields are required'}), 400

    # Hash the password
    hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')

    # Insert user into database
    try:
        cursor = mysql.connection.cursor()
        cursor.execute('INSERT INTO users (username, email, password,roles) VALUES (%s, %s, %s,%s)',
                       (username, email, hashed_password,roles,))
        mysql.connection.commit()
        cursor.close()
        return jsonify({'message': 'User registered successfully'}), 201
    except Exception as e:
        return jsonify({'message': 'User registration failed', 'error': str(e)}), 500

# Route for user login
@app.route('/login', methods=['POST'])
def login():
    data = request.json
    email = data.get('email')
    password = data.get('password')

    # Check for missing fields
    if not email or not password:
        return jsonify({'message': 'Both email and password are required'}), 400

    # Fetch user from database
    cursor = mysql.connection.cursor()
    cursor.execute('SELECT * FROM users WHERE email = %s', (email,))
    user = cursor.fetchone()
    cursor.close()

    if user and bcrypt.check_password_hash(user[3], password):
        return jsonify({'message': 'Login successful', 'user': {'id': user[0], 'username': user[1], 'email': user[2],'roles':user[4]}}), 200
    else:
        return jsonify({'message': 'Invalid email or password'}), 401
@app.route('/save_profile', methods=['POST'])
def save_profile():
    data = request.form  # Use `request.form` for form data
    user_id = data.get('user_id')
    full_name = data.get('full_name')
    phone_number = data.get('phone_number')
    address = data.get('address')
    education = data.get('education')
    work_experience = data.get('work_experience')
    skills = data.get('skills')
    linkedin_profile = data.get('linkedin_profile')
    github_profile = data.get('github_profile')

    # Check for missing fields
    if not user_id or not full_name or not education or not work_experience:
        return jsonify({'message': 'User ID, full name, education, and work experience are required'}), 400

    # Check if the user exists in the users table
    cursor = mysql.connection.cursor()
    cursor.execute('SELECT id FROM users WHERE id = %s', (user_id,))
    user_exists = cursor.fetchone()

    if not user_exists:
        return jsonify({'message': 'User ID does not exist in the users table'}), 400

    # Check if an image is uploaded
    profile_image = request.files.get('profile_image')  # Assuming the field name in the form is 'profile_image'
    image_data = None
    
    if profile_image and allowed_file(profile_image.filename):
        # Read the image data as binary
        image_data = profile_image.read()

    # Check if the user profile already exists
    cursor.execute('SELECT * FROM user_profiles WHERE user_id = %s', (user_id,))
    existing_profile = cursor.fetchone()

    if existing_profile:
        # If an image is not uploaded, keep the existing image in the database
        if not image_data:
            image_data = existing_profile[10]  # Access the 12th column directly

        # Update the existing profile
        try:
            cursor.execute('''UPDATE user_profiles
                               SET full_name = %s, phone_number = %s, address = %s, education = %s,
                                   work_experience = %s, skills = %s, linkedin_profile = %s, github_profile = %s,
                                   profile_image = %s
                               WHERE user_id = %s''',
                           (full_name, phone_number, address, education, work_experience, skills, linkedin_profile, github_profile, image_data, user_id))
            mysql.connection.commit()
            cursor.close()
            return jsonify({'message': 'Profile updated successfully'}), 200
        except Exception as e:
            return jsonify({'message': 'Failed to update profile', 'error': str(e)}), 500
    else:
        # Create a new profile
        try:
            cursor.execute('''INSERT INTO user_profiles
                               (user_id, full_name, phone_number, address, education, work_experience, skills, linkedin_profile, github_profile, profile_image)
                               VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)''',
                           (user_id, full_name, phone_number, address, education, work_experience, skills, linkedin_profile, github_profile, image_data))
            mysql.connection.commit()
            cursor.close()
            return jsonify({'message': 'Profile saved successfully'}), 201
        except Exception as e:
            return jsonify({'message': 'Failed to save profile', 'error': str(e)}), 500


@app.route('/get_profile/<int:user_id>', methods=['GET'])
def get_profile(user_id):
    cursor = mysql.connection.cursor()
    cursor.execute('SELECT * FROM user_profiles WHERE user_id = %s', (user_id,))
    profile = cursor.fetchone()
    cursor.close()

    if profile:
        # Assuming the profile image is at index 10 (adjust if necessary)
        profile_image = profile[10]  # Index 10 for profile_image (binary data)
        profile_image_base64 = None

        # Check if the profile_image is not None and is of the correct type (bytes)
        if profile_image:
            if isinstance(profile_image, bytes):  # If it's a byte-like object
                profile_image_base64 = base64.b64encode(profile_image).decode('utf-8')
            else:
                print(f"Unexpected profile_image type: {type(profile_image)}")
                # You can either send a default image or handle the scenario as needed
                profile_image_base64 = get_default_image_base64()

        return jsonify({
            'user_id': profile[1],
            'full_name': profile[2],
            'phone_number': profile[3],
            'address': profile[4],
            'education': profile[5],
            'work_experience': profile[6],
            'skills': profile[7],
            'linkedin_profile': profile[8],
            'github_profile': profile[9],
            'profile_image': profile_image_base64  # Send the image as base64 string
        }), 200
    else:
        return jsonify({'message': 'Profile not found'}), 404


def get_default_image_base64():
    """
    Helper function to return a default image as base64 if no profile image is available
    """
    # Specify a default image path (relative or absolute)
    default_image_path = 'path/to/default_image.png'
    
    try:
        with open(default_image_path, 'rb') as image_file:
            default_image_data = image_file.read()
            return base64.b64encode(default_image_data).decode('utf-8')
    except Exception as e:
        print(f"Error loading default image: {e}")
        return None








# Function to generate and fill DOCX with user data


def generate_profile_doc(profile_data):
    # Load the template DOCX file
    template_path = r'C:\Users\Win\Downloads\templete1.docx'  # Use raw string for Windows path
    document = Document(template_path)

    # Replace placeholders with actual data
    for paragraph in document.paragraphs:
        for key, value in profile_data.items():
            placeholder = '{' + key + '}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Save the modified DOCX
    output_path = rf'C:\Users\Win\Downloads\profile_{profile_data["user_id"]}.docx'
    document.save(output_path)
    
    return output_path






@app.route('/get_profile_doc/<int:user_id>/<int:template_id>', methods=['GET'])
def get_profile_doc(user_id, template_id):
    # Fetch user profile from the database
    cursor = mysql.connection.cursor()
    cursor.execute('SELECT * FROM user_profiles WHERE user_id = %s', (user_id,))
    profile = cursor.fetchone()

    if not profile:
        cursor.close()
        return jsonify({'message': 'Profile not found'}), 404

    # Fetch the template from the database
    cursor.execute('SELECT template_data FROM templates WHERE template_id = %s', (template_id,))
    template_row = cursor.fetchone()
    cursor.close()

    if not template_row:
        return jsonify({'message': 'Template not found'}), 404

    # Map database row to dictionary
    profile_data = {
        'user_id': profile[1],
        'full_name': profile[2],
        'phone_number': profile[3],
        'address': profile[4],
        'education': profile[5],
        'work_experience': profile[6],
        'skills': profile[7],
        'linkedin_profile': profile[8],
        'github_profile': profile[9]
    }

    # Load the template DOCX data from the database
    template_data = template_row[0]
    with open("temp_template.docx", "wb") as temp_file:
        temp_file.write(template_data)

    # Load and modify the document
    document = Document("temp_template.docx")
    for paragraph in document.paragraphs:
        for key, value in profile_data.items():
            placeholder = '{' + key + '}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Save the generated DOCX file
    output_path = f'profile_{profile_data["user_id"]}.docx'
    document.save(output_path)

    # Send the DOCX file to the frontend
    return send_file(output_path, as_attachment=True)




@app.route('/upload_template', methods=['POST'])
def upload_template():
    # Get the DOCX file and image file from the request
    file = request.files.get('file')  # DOCX file
    image = request.files.get('image')  # Image file
    template_name = request.form.get('template_name')

    if not file or not image or not template_name:
        return jsonify({'message': 'Template file, image, and name are required'}), 400

    # Read the DOCX file data
    template_data = file.read()
    template_image = image.read()  # Read the image file data

    try:
        cursor = mysql.connection.cursor()
        # Insert the template and image into the database
        cursor.execute('''INSERT INTO templates (template_name, template_data, template_image)
                          VALUES (%s, %s, %s)''', 
                       (template_name, template_data, template_image))
        mysql.connection.commit()
        cursor.close()
        return jsonify({'message': 'Template and image uploaded successfully'}), 201
    except Exception as e:
        return jsonify({'message': 'Failed to upload template and image', 'error': str(e)}), 500




# Route to get all templatesimport base64


@app.route('/get_templates', methods=['GET'])
def get_templates():
    cursor = mysql.connection.cursor()
    cursor.execute('SELECT template_id, template_name, template_image FROM templates')
    templates = cursor.fetchall()
    cursor.close()

    if templates:
        result = []
        for template in templates:
            template_id, template_name, template_image = template

            # If there's an image, base64 encode it
            if template_image:
                # Convert the image to base64
                image_data = base64.b64encode(template_image).decode('utf-8')
            else:
                image_data = None

            result.append({
                'template_id': template_id,
                'template_name': template_name,
                'template_image': image_data  # Send the base64 encoded image
            })

        return jsonify({'templates': result}), 200
    else:
        return jsonify({'message': 'No templates found'}), 404


@app.route('/delete_template/<int:template_id>', methods=['DELETE'])
def delete_template(template_id):
    try:
        cursor = mysql.connection.cursor()
        # Delete the template from the database based on the template_id
        cursor.execute('DELETE FROM templates WHERE template_id = %s', (template_id,))
        mysql.connection.commit()
        cursor.close()

        # Check if a template was deleted
        if cursor.rowcount == 0:
            return jsonify({'message': 'Template not found'}), 404
        else:
            return jsonify({'message': 'Template deleted successfully'}), 200
    except Exception as e:
        return jsonify({'message': 'Failed to delete template', 'error': str(e)}), 500


# Run the Flask app
if __name__ == '__main__':
    app.run(debug=True)
